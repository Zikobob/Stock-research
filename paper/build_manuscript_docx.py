"""
Build an editable Word (.docx) of the paper in a classic student-manuscript
style (Times New Roman 12, double-spaced, centered bold headings, page numbers
top-right, an author-affiliation footer, and a formal title page).

Pipeline:
  1. Create a reference-doc that defines the styles + header/footer.
  2. Convert research_paper.md -> body via pandoc using that reference-doc.
  3. Post-process with python-docx to prepend the title page.

Run:  python paper/build_manuscript_docx.py
Output: paper/research_paper.docx
"""

from __future__ import annotations

import os
import re

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(HERE, "research_paper.md")
REF = os.path.join(HERE, "_reference.docx")
OUT = os.path.join(HERE, "research_paper.docx")

# ----- Author / affiliation (edit these) ----------------------------------- #
AUTHOR = "Aarav Vaidha"
MENTOR = "Guidance from [Mentor / Advisor]"          # edit or delete
AFFILIATION = "MOT Charter High School, 1275 Cedar Lane Rd, Middletown, DE 19709, USA."
EMAIL = "aaravvaidha69@gmail.com"
TITLE = ("Cross-Sector Correlation Regimes and Their Impact on the Predictive "
         "Accuracy of Equity Price Forecasting Models")
KEYWORDS = ("cross-sector correlation, realized volatility, regime analysis, "
            "out-of-sample forecasting, return predictability, volatility "
            "confound, Newey–West standard errors")

TNR = "Times New Roman"


def _set_font(style, size=12, bold=False):
    style.font.name = size and TNR
    style.font.size = Pt(size)
    style.font.bold = bold
    # Ensure the east-asian/complex slots also use TNR.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), TNR)


def _page_number_field(paragraph):
    run = paragraph.add_run()
    run.font.name = TNR
    run.font.size = Pt(12)
    for typ, extra in (("begin", None), ("text", "PAGE"), ("end", None)):
        if typ == "text":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), typ)
        run._r.append(el)


def _top_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "6"),
                 ("w:color", "auto")):
        top.set(qn(k), v)
    pbdr.append(top)
    pPr.append(pbdr)


def build_reference_doc() -> None:
    doc = Document()

    # --- Normal: TNR 12, double spacing, first-line indent -----------------
    normal = doc.styles["Normal"]
    _set_font(normal, 12)
    pf = normal.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = Inches(0.5)

    # --- Headings: TNR 12 bold black; H1/H2 centered, H3+ left -------------
    for lvl in range(1, 5):
        try:
            h = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        _set_font(h, 12, bold=True)
        h.font.color.rgb = None  # inherit black
        hpf = h.paragraph_format
        hpf.first_line_indent = Inches(0)
        hpf.space_before = Pt(12)
        hpf.space_after = Pt(6)
        hpf.line_spacing = 2.0
        hpf.keep_with_next = True
        hpf.alignment = (WD_ALIGN_PARAGRAPH.CENTER if lvl <= 2
                         else WD_ALIGN_PARAGRAPH.LEFT)
        # Force black text (Word default heading colour is blue).
        rpr = h.element.get_or_add_rPr()
        color = rpr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rpr.append(color)
        color.set(qn("w:val"), "000000")

    sec = doc.sections[0]
    sec.different_first_page_header_footer = False

    # --- Header: page number, right-aligned --------------------------------
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _page_number_field(hp)

    # --- Footer: rule + affiliation footnote line --------------------------
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.first_line_indent = Inches(0)
    _top_border(fp)
    sup = fp.add_run("1")
    sup.font.name = TNR
    sup.font.size = Pt(11)
    sup.font.superscript = True
    rest = fp.add_run(f"{AUTHOR}, {AFFILIATION} {EMAIL}")
    rest.font.name = TNR
    rest.font.size = Pt(11)

    doc.save(REF)


def stripped_markdown() -> str:
    with open(MD, encoding="utf-8") as f:
        text = f.read()
    # Drop the leading H1 title + subtitle; the title page replaces them.
    idx = text.index("## Abstract")
    return text[idx:]


def add_title_page(doc: Document) -> None:
    """Insert the manuscript title page before the first body paragraph."""
    first = doc.paragraphs[0]

    def new(before, text="", *, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
            size=12, blank_after=0, superscript_tail=None):
        p = before.insert_paragraph_before()
        p.alignment = align
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        if text:
            r = p.add_run(text)
            r.font.name = TNR
            r.font.size = Pt(size)
            r.font.bold = bold
        if superscript_tail:
            s = p.add_run(superscript_tail)
            s.font.name = TNR
            s.font.size = Pt(size)
            s.font.superscript = True
        return p

    # Title (top-left)
    new(first, TITLE, align=WD_ALIGN_PARAGRAPH.LEFT)
    for _ in range(8):
        new(first)  # vertical gap
    # Centered block
    new(first, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    new(first, AUTHOR, align=WD_ALIGN_PARAGRAPH.CENTER, superscript_tail="1")
    if MENTOR:
        new(first, MENTOR, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    for _ in range(3):
        new(first)
    # Keywords
    kw = new(first, "", align=WD_ALIGN_PARAGRAPH.LEFT)
    b = kw.add_run("Keywords: ")
    b.font.name = TNR
    b.font.size = Pt(12)
    b.font.bold = True
    k = kw.add_run(KEYWORDS)
    k.font.name = TNR
    k.font.size = Pt(12)
    # Page break to end the title page
    pb = new(first)
    pb.add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    build_reference_doc()

    tmp_md = os.path.join(HERE, "_body.md")
    with open(tmp_md, "w", encoding="utf-8") as f:
        f.write(stripped_markdown())

    body = os.path.join(HERE, "_body.docx")
    pypandoc.convert_file(
        tmp_md, "docx", outputfile=body,
        extra_args=["--reference-doc", REF, "--standalone",
                    "--resource-path", HERE],
    )

    doc = Document(body)
    add_title_page(doc)
    doc.save(OUT)

    for f in (tmp_md, body, REF):
        try:
            os.remove(f)
        except OSError:
            pass
    print(f"Wrote {OUT} ({os.path.getsize(OUT)/1024:.0f} KB)")


if __name__ == "__main__":
    main()
